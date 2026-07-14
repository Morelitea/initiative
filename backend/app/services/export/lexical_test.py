"""Unit tests for the Lexical editor-state converters."""

import io
import zipfile

import pytest

from app.services.export.lexical import (
    blocks_from_editor_state,
    render_docx,
    render_markdown,
)

pytestmark = pytest.mark.unit

GUILD = 7


def _text(text, fmt=0):
    return {"type": "text", "text": text, "format": fmt}


def _state(children):
    return {"root": {"type": "root", "children": children}}


FIXTURE = _state(
    [
        {"type": "heading", "tag": "h2", "children": [_text("Overview")]},
        {
            "type": "paragraph",
            "children": [
                _text("plain "),
                _text("bold", 1),
                _text(" italic", 2),
                _text(" code", 16),
                {
                    "type": "link",
                    "url": "https://example.com",
                    "children": [_text("a link")],
                },
                {"type": "mention", "text": "@alice", "mentionName": "alice"},
            ],
        },
        {"type": "quote", "children": [_text("wisdom")]},
        {
            "type": "code",
            "language": "python",
            "children": [_text("x = 1"), {"type": "linebreak"}, _text("y = 2")],
        },
        {
            "type": "list",
            "listType": "check",
            "children": [
                {"type": "listitem", "checked": True, "children": [_text("done")]},
                {
                    "type": "listitem",
                    "checked": False,
                    "children": [
                        _text("todo"),
                        {
                            "type": "list",
                            "listType": "bullet",
                            "children": [
                                {"type": "listitem", "children": [_text("nested")]}
                            ],
                        },
                    ],
                },
            ],
        },
        {"type": "image", "src": f"/uploads/{GUILD}/img-abc.png", "altText": "our pic"},
        {"type": "image", "src": "https://elsewhere.example/x.png", "altText": "ext"},
        {"type": "horizontalrule"},
        {"type": "youtube", "videoID": "dQw4w9WgXcQ"},
        {
            "type": "unknown-future-node",
            "children": [{"type": "paragraph", "children": [_text("still exported")]}],
        },
    ]
)


def test_parser_blocks_and_assets():
    blocks, assets = blocks_from_editor_state(FIXTURE, guild_id=GUILD)
    types = [b["type"] for b in blocks]
    assert types == [
        "heading",
        "paragraph",
        "quote",
        "code",
        "list",
        "image",
        "image",
        "hr",
        "paragraph",  # youtube degrades to a link paragraph
        "paragraph",  # unknown container's child
    ]
    # Same-guild upload collected as an asset; external stays a URL.
    assert assets == [{"key": "img-abc.png", "name": "img-abc.png"}]
    assert blocks[5]["asset"] == "img-abc.png"
    assert blocks[6]["asset"] is None
    assert blocks[6]["url"] == "https://elsewhere.example/x.png"
    # Format bitmask decoded; mention degraded to its text.
    para = blocks[1]["runs"]
    assert {"text": "bold", "bold": True} in para
    assert {"text": " code", "code": True} in para
    assert {"text": "a link", "link": "https://example.com"} in para
    assert {"text": "@alice"} in para
    # Code block joins its lines.
    assert blocks[3]["text"] == "x = 1\ny = 2"
    # Checklist with nesting.
    items = blocks[4]["items"]
    assert items[0]["checked"] is True
    assert items[1]["children"][0]["items"][0]["runs"] == [{"text": "nested"}]


def test_markdown_zips_when_assets_present():
    blocks, assets = blocks_from_editor_state(FIXTURE, guild_id=GUILD)
    data = {"title": "Notes", "stem": "notes", "blocks": blocks, "assets": assets}
    content, content_type, filename = render_markdown(
        data, lambda key: b"png-bytes-" + key.encode()
    )
    assert content_type == "application/zip"
    assert filename == "notes.zip"
    archive = zipfile.ZipFile(io.BytesIO(content))
    assert set(archive.namelist()) == {"notes.md", "assets/img-abc.png"}
    assert archive.read("assets/img-abc.png") == b"png-bytes-img-abc.png"
    md = archive.read("notes.md").decode("utf-8")
    assert "# Notes" in md
    assert "### Overview" in md  # content h2 shifted below the title
    assert "**bold**" in md and "` code`" in md
    assert "[a link](<https://example.com>)" in md
    assert "> wisdom" in md
    assert "```python" in md
    assert "- [x] done" in md and "- [ ] todo" in md
    assert "  - nested" in md
    assert "![our pic](assets/img-abc.png)" in md
    assert "![ext](<https://elsewhere.example/x.png>)" in md
    assert "https://www.youtube.com/watch?v=dQw4w9WgXcQ" in md
    assert "still exported" in md


def test_markdown_table_escapes_pipes():
    state = _state(
        [
            {
                "type": "table",
                "children": [
                    {
                        "type": "tablerow",
                        "children": [
                            {
                                "type": "tablecell",
                                "children": [
                                    {
                                        "type": "paragraph",
                                        "children": [_text("a | b")],
                                    }
                                ],
                            }
                        ],
                    }
                ],
            }
        ]
    )
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    content, _, _ = render_markdown(
        {"title": "", "blocks": blocks, "assets": assets}, lambda key: b""
    )
    md = content.decode("utf-8")
    assert "| a \\| b |" in md  # a raw pipe would split the column


def test_asset_names_deduped_on_sanitize_collision():
    """Two distinct storage keys can sanitize to one archive name — each must
    keep its own entry or one image silently replaces the other."""
    state = _state(
        [
            {"type": "image", "src": f"/uploads/{GUILD}/a img.png", "altText": "one"},
            {"type": "image", "src": f"/uploads/{GUILD}/a_img.png", "altText": "two"},
        ]
    )
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    names = [a["name"] for a in assets]
    assert len(assets) == 2
    assert len(set(names)) == 2  # unique archive names
    # Blocks reference their own (deduped) names.
    assert {b["asset"] for b in blocks} == set(names)
    content, content_type, _ = render_markdown(
        {"title": "", "stem": "d", "blocks": blocks, "assets": assets},
        lambda key: key.encode(),
    )
    archive = zipfile.ZipFile(io.BytesIO(content))
    stored = {n: archive.read(n) for n in archive.namelist() if n.startswith("assets/")}
    assert len(stored) == 2
    assert set(stored.values()) == {b"a img.png", b"a_img.png"}  # both survive


def test_markdown_zip_survives_missing_asset():
    """An image gone from storage skips its archive entry; the export (and
    the other assets) still ship."""
    state = _state(
        [
            {"type": "image", "src": f"/uploads/{GUILD}/gone.png", "altText": "gone"},
            {"type": "image", "src": f"/uploads/{GUILD}/here.png", "altText": "here"},
        ]
    )
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)

    def read(key: str) -> bytes:
        if key == "gone.png":
            raise FileNotFoundError(key)
        return b"bytes"

    content, content_type, _ = render_markdown(
        {"title": "", "stem": "d", "blocks": blocks, "assets": assets}, read
    )
    assert content_type == "application/zip"
    archive = zipfile.ZipFile(io.BytesIO(content))
    assert "assets/here.png" in archive.namelist()
    assert "assets/gone.png" not in archive.namelist()


def test_markdown_inline_code_with_backticks():
    state = _state(
        [
            {
                "type": "paragraph",
                "children": [_text("use `backticks` here", 16)],  # code format
            }
        ]
    )
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    content, _, _ = render_markdown(
        {"title": "", "blocks": blocks, "assets": assets}, lambda key: b""
    )
    # Double-backtick spaced form keeps the inner backticks literal.
    assert "`` use `backticks` here ``" in content.decode("utf-8")


def test_markdown_plain_without_assets():
    state = _state([{"type": "paragraph", "children": [_text("hello")]}])
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    content, content_type, filename = render_markdown(
        {"title": "T", "stem": "t", "blocks": blocks, "assets": assets},
        lambda key: b"",
    )
    assert content_type.startswith("text/markdown")
    assert filename is None
    assert "hello" in content.decode("utf-8")


def test_docx_renders_and_embeds_image():
    import struct
    import zlib

    def make_png():
        def chunk(tag, data):
            c = tag + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b"")
        )

    import docx

    blocks, assets = blocks_from_editor_state(FIXTURE, guild_id=GUILD)
    data = {"title": "Notes", "stem": "notes", "blocks": blocks, "assets": assets}
    content = render_docx(data, lambda key: make_png())
    assert content.startswith(b"PK")
    document = docx.Document(io.BytesIO(content))
    texts = [p.text for p in document.paragraphs]
    assert "Notes" in texts  # title
    assert any("wisdom" in t for t in texts)
    assert any("still exported" in t for t in texts)
    assert any("☑ done" in t for t in texts)
    assert len(document.inline_shapes) == 1  # the embedded upload image


def test_image_resize_carried_and_honored_in_docx():
    """The editor's resize (width px on the node) must survive into the
    block and set the DOCX picture width — not the fixed 6-inch default."""
    import struct
    import zlib

    import docx
    from docx.shared import Inches

    def make_png():
        def chunk(tag, data):
            c = tag + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b"")
        )

    state = _state(
        [
            {
                "type": "image",
                "src": f"/uploads/{GUILD}/pic.png",
                "altText": "sized",
                "width": 192,  # 2 inches at 96 dpi
                "height": 120,
            },
            {
                "type": "image",
                "src": f"/uploads/{GUILD}/pic.png",
                "altText": "untouched",
                "width": "inherit",  # editor default: no resize stored
            },
        ]
    )
    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    assert blocks[0]["width"] == 192  # width wins; stale height not carried
    assert "height" not in blocks[0]
    assert "width" not in blocks[1]  # "inherit" ignored

    content = render_docx(
        {"title": "", "blocks": blocks, "assets": assets}, lambda key: make_png()
    )
    document = docx.Document(io.BytesIO(content))
    shapes = document.inline_shapes
    assert len(shapes) == 2
    assert shapes[0].width == Inches(2)  # resized: 192px / 96dpi
    assert shapes[1].width == Inches(6)  # untouched: the default cap


def test_docx_degrades_unreadable_image_to_alt_text():
    state = _state(
        [{"type": "image", "src": f"/uploads/{GUILD}/bad.png", "altText": "broken"}]
    )
    import docx

    blocks, assets = blocks_from_editor_state(state, guild_id=GUILD)
    content = render_docx(
        {"title": "", "blocks": blocks, "assets": assets},
        lambda key: b"not-an-image",
    )
    document = docx.Document(io.BytesIO(content))
    assert any("broken" in p.text for p in document.paragraphs)
