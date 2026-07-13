"""Lexical editor-state conversion for document exports (md / pdf / docx).

One tree walk (``blocks_from_editor_state``) turns the serialized editor
state into a JSON-safe intermediate — a flat list of typed blocks with
inline "runs" — that all three renderers consume: the Markdown emitter
(zipped with an ``assets/`` folder when images are referenced), the DOCX
emitter (python-docx, images embedded), and the ``document`` Typst template
(images staged into the compile root by the backend).

Degradation contract: any unknown node with ``text`` renders as text (this
covers mentions, wikilinks, hashtags, emojis, keywords — all TextNode
subclasses); unknown containers recurse into their children; embeds
(YouTube/Tweet) degrade to a link; anything else is dropped silently. A new
editor node can never break an export, it just exports as its text.

Images: only same-guild uploads (``/uploads/{guild_id}/…``) are collected as
assets and embedded — an external URL is never fetched (no SSRF surface, per
the export design); it renders as a plain link instead.
"""

from __future__ import annotations

import io
import zipfile
from typing import Any, Callable

from app.services.platform.csv_export import safe_filename_component

# Lexical TextNode format bitmask.
_BOLD = 1
_ITALIC = 2
_STRIKE = 4
_UNDERLINE = 8
_CODE = 16

ReadBlob = Callable[[str], bytes]

# ---------------------------------------------------------------------------
# Parse: editor state -> blocks + referenced assets
# ---------------------------------------------------------------------------


def blocks_from_editor_state(
    content: dict, *, guild_id: int
) -> tuple[list[dict], list[dict]]:
    """Walk the serialized editor state into (blocks, assets).

    ``assets`` lists the same-guild upload blobs the document references:
    ``{"key": <storage key>, "name": <archive-safe file name>}``.
    """
    parser = _Parser(guild_id)
    root = (content or {}).get("root") or {}
    parser.walk_children(root.get("children") or [])
    parser.flush_paragraph()
    return parser.blocks, list(parser.assets.values())


class _Parser:
    def __init__(self, guild_id: int) -> None:
        self.guild_id = guild_id
        self.blocks: list[dict] = []
        self.assets: dict[str, dict] = {}
        self._pending_runs: list[dict] = []

    # -- inline ------------------------------------------------------------

    def collect_runs(self, children: list, link: str | None = None) -> list[dict]:
        runs: list[dict] = []
        for node in children or []:
            ntype = node.get("type")
            if ntype in ("link", "autolink"):
                runs.extend(
                    self.collect_runs(
                        node.get("children") or [], link=node.get("url") or link
                    )
                )
                continue
            if ntype == "linebreak":
                runs.append({"text": "\n"})
                continue
            if ntype == "image":
                # An inline image splits the paragraph: emit what we have,
                # then the image as its own block.
                self._pending_runs.extend(runs)
                runs = []
                self.flush_paragraph()
                self.image_block(node)
                continue
            text = node.get("text")
            if isinstance(text, str) and text:
                fmt = int(node.get("format") or 0)
                run: dict[str, Any] = {"text": text}
                if fmt & _BOLD:
                    run["bold"] = True
                if fmt & _ITALIC:
                    run["italic"] = True
                if fmt & _STRIKE:
                    run["strike"] = True
                if fmt & _UNDERLINE:
                    run["underline"] = True
                if fmt & _CODE:
                    run["code"] = True
                if link:
                    run["link"] = link
                runs.append(run)
                continue
            # Unknown node: recurse if it has children, else drop.
            if node.get("children"):
                runs.extend(self.collect_runs(node["children"], link=link))
        return runs

    def flush_paragraph(self) -> None:
        if self._pending_runs:
            self.blocks.append({"type": "paragraph", "runs": self._pending_runs})
            self._pending_runs = []

    # -- blocks ------------------------------------------------------------

    def walk_children(self, children: list) -> None:
        for node in children or []:
            self.block(node)

    def block(self, node: dict) -> None:
        ntype = node.get("type")
        children = node.get("children") or []
        if ntype == "paragraph":
            runs = self.collect_runs(children)
            runs = self._pending_runs + runs
            self._pending_runs = []
            if runs:
                self.blocks.append({"type": "paragraph", "runs": runs})
        elif ntype == "heading":
            tag = str(node.get("tag") or "h1")
            level = int(tag[1]) if len(tag) == 2 and tag[1].isdigit() else 1
            self.blocks.append(
                {"type": "heading", "level": level, "runs": self.collect_runs(children)}
            )
            self.flush_paragraph()
        elif ntype == "quote":
            self.blocks.append({"type": "quote", "runs": self.collect_runs(children)})
            self.flush_paragraph()
        elif ntype == "code":
            self.blocks.append(
                {
                    "type": "code",
                    "language": str(node.get("language") or ""),
                    "text": _code_text(children),
                }
            )
        elif ntype == "list":
            self.blocks.append(self.list_block(node))
        elif ntype == "horizontalrule":
            self.blocks.append({"type": "hr"})
        elif ntype == "table":
            self.blocks.append(self.table_block(node))
        elif ntype == "image":
            self.image_block(node)
        elif ntype in ("youtube", "tweet"):
            url = _embed_url(node)
            if url:
                self.blocks.append(
                    {"type": "paragraph", "runs": [{"text": url, "link": url}]}
                )
        elif children:
            # Unknown container (layout containers/items, future nodes):
            # its children still export.
            self.walk_children(children)

    def list_block(self, node: dict) -> dict:
        list_type = node.get("listType")  # "bullet" | "number" | "check"
        items: list[dict] = []
        for child in node.get("children") or []:
            if child.get("type") != "listitem":
                continue
            nested = [
                gc for gc in (child.get("children") or []) if gc.get("type") == "list"
            ]
            inline = [
                gc for gc in (child.get("children") or []) if gc.get("type") != "list"
            ]
            item: dict[str, Any] = {"runs": self.collect_runs(inline)}
            if list_type == "check":
                item["checked"] = bool(child.get("checked"))
            if nested:
                item["children"] = [self.list_block(n) for n in nested]
            items.append(item)
        return {
            "type": "list",
            "ordered": list_type == "number",
            "checklist": list_type == "check",
            "items": items,
        }

    def table_block(self, node: dict) -> dict:
        rows: list[list[list[dict]]] = []
        for row in node.get("children") or []:
            if row.get("type") != "tablerow":
                continue
            cells: list[list[dict]] = []
            for cell in row.get("children") or []:
                if cell.get("type") != "tablecell":
                    continue
                cell_runs: list[dict] = []
                for para in cell.get("children") or []:
                    if cell_runs:
                        cell_runs.append({"text": "\n"})
                    cell_runs.extend(self.collect_runs(para.get("children") or []))
                cells.append(cell_runs)
            if cells:
                rows.append(cells)
        return {"type": "table", "rows": rows}

    def image_block(self, node: dict) -> None:
        src = str(node.get("src") or "")
        alt = str(node.get("altText") or "")
        prefix = f"/uploads/{self.guild_id}/"
        if src.startswith(prefix):
            key = src.removeprefix(prefix).split("?")[0]
            existing = self.assets.get(key)
            if existing is None:
                # Distinct keys can sanitize to the same name ("a img.png" /
                # "a_img.png") — suffix until unique so one archive entry
                # can't silently overwrite another.
                name = safe_filename_component(key)
                taken = {a["name"] for a in self.assets.values()}
                if name in taken:
                    stem, dot, ext = name.rpartition(".")
                    base = stem if dot else name
                    n = 2
                    while name in taken:
                        name = f"{base}-{n}{dot}{ext}" if dot else f"{base}-{n}"
                        n += 1
                existing = {"key": key, "name": name}
                self.assets[key] = existing
            block = {"type": "image", "asset": existing["name"], "alt": alt}
            block.update(_image_size(node))
            self.blocks.append(block)
        elif src:
            # External (or cross-guild) image: never fetched — link only.
            self.blocks.append({"type": "image", "asset": None, "url": src, "alt": alt})


def _image_size(node: dict) -> dict:
    """The editor stores a resize as width/height px on the node ("inherit"
    when untouched). Width alone is carried — the renderers derive height
    from the aspect ratio, so a stale stored height can't distort."""
    for field in ("width", "height"):
        value = node.get(field)
        if (
            isinstance(value, (int, float))
            and not isinstance(value, bool)
            and value > 0
        ):
            return {field: int(value)}
    return {}


def _code_text(children: list) -> str:
    parts: list[str] = []
    for node in children or []:
        if node.get("type") == "linebreak":
            parts.append("\n")
        else:
            text = node.get("text")
            if isinstance(text, str):
                parts.append(text)
    return "".join(parts)


def _embed_url(node: dict) -> str | None:
    if node.get("type") == "youtube" and node.get("videoID"):
        return f"https://www.youtube.com/watch?v={node['videoID']}"
    if node.get("type") == "tweet" and node.get("tweetID"):
        return f"https://x.com/i/status/{node['tweetID']}"
    return None


# ---------------------------------------------------------------------------
# Markdown
# ---------------------------------------------------------------------------


def render_markdown(data: dict, read_blob: ReadBlob) -> tuple[bytes, str, str | None]:
    """Emit Markdown. With referenced assets: a zip of ``{stem}.md`` +
    ``assets/``, image refs rewritten to the relative paths. Returns
    (content, content_type, filename_override)."""
    stem = str(data.get("stem") or "document")
    md = _markdown_text(data)
    assets = data.get("assets") or []
    if not assets:
        return md.encode("utf-8"), "text/markdown; charset=utf-8", None

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{stem}.md", md)
        for asset in assets:
            archive.writestr(f"assets/{asset['name']}", read_blob(asset["key"]))
    return buffer.getvalue(), "application/zip", f"{stem}.zip"


def _markdown_text(data: dict) -> str:
    lines: list[str] = []
    title = str(data.get("title") or "").strip()
    if title:
        lines.append(f"# {title}")
        lines.append("")
    for block in data.get("blocks") or []:
        lines.extend(_md_block(block))
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def _md_block(block: dict, indent: str = "") -> list[str]:
    btype = block.get("type")
    if btype == "heading":
        # Document title is H1; shift content headings down one level.
        level = min(int(block.get("level") or 1) + 1, 6)
        return [f"{'#' * level} {_md_runs(block.get('runs') or [])}"]
    if btype == "quote":
        return [f"> {_md_runs(block.get('runs') or [])}"]
    if btype == "code":
        return [f"```{block.get('language') or ''}", block.get("text") or "", "```"]
    if btype == "hr":
        return ["---"]
    if btype == "image":
        alt = _md_escape(block.get("alt") or "")
        if block.get("asset"):
            return [f"![{alt}](assets/{block['asset']})"]
        return [f"![{alt}](<{block.get('url') or ''}>)"]
    if btype == "list":
        return _md_list(block, indent)
    if btype == "table":
        return _md_table(block)
    return [f"{indent}{_md_runs(block.get('runs') or [])}"]


def _md_list(block: dict, indent: str = "") -> list[str]:
    lines: list[str] = []
    for index, item in enumerate(block.get("items") or [], start=1):
        if block.get("checklist"):
            box = "x" if item.get("checked") else " "
            marker = f"- [{box}]"
        elif block.get("ordered"):
            marker = f"{index}."
        else:
            marker = "-"
        lines.append(f"{indent}{marker} {_md_runs(item.get('runs') or [])}")
        for nested in item.get("children") or []:
            lines.extend(_md_list(nested, indent + "  "))
    return lines


def _md_table(block: dict) -> list[str]:
    rows = block.get("rows") or []
    if not rows:
        return []
    width = max(len(r) for r in rows)

    def cells(row: list) -> str:
        padded = list(row) + [[]] * (width - len(row))
        # Escape pipes AFTER rendering (a raw | in cell text would split the
        # column); _md_escape would also mangle the link syntax just emitted.
        rendered = (_md_runs(c).replace("\n", " ").replace("|", "\\|") for c in padded)
        return "| " + " | ".join(rendered) + " |"

    lines = [cells(rows[0]), "|" + "|".join(" --- " for _ in range(width)) + "|"]
    lines.extend(cells(row) for row in rows[1:])
    return lines


def _md_escape(text: str) -> str:
    return text.replace("|", "\\|").replace("[", "\\[").replace("]", "\\]")


def _md_runs(runs: list[dict]) -> str:
    parts: list[str] = []
    for run in runs or []:
        text = run.get("text") or ""
        if text == "\n":
            parts.append("  \n")
            continue
        if run.get("code"):
            text = f"`{text}`"
        else:
            if run.get("bold") and run.get("italic"):
                text = f"***{text}***"
            elif run.get("bold"):
                text = f"**{text}**"
            elif run.get("italic"):
                text = f"*{text}*"
            if run.get("strike"):
                text = f"~~{text}~~"
        if run.get("link"):
            text = f"[{text}](<{run['link']}>)"
        parts.append(text)
    return "".join(parts)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def render_docx(data: dict, read_blob: ReadBlob) -> bytes:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.shared import Inches, Pt

    document = docx.Document()
    title = str(data.get("title") or "").strip()
    if title:
        document.add_heading(title, level=0)

    def add_runs(paragraph, runs: list[dict]) -> None:
        for run in runs or []:
            text = run.get("text") or ""
            if text == "\n":
                paragraph.add_run().add_break(WD_BREAK.LINE)
                continue
            r = paragraph.add_run(text)
            r.bold = bool(run.get("bold")) or None
            r.italic = bool(run.get("italic")) or None
            r.font.strike = bool(run.get("strike")) or None
            r.underline = bool(run.get("underline")) or None
            if run.get("code"):
                r.font.name = "Consolas"
            if run.get("link"):
                # python-docx has no first-class hyperlink on runs; keep the
                # text and append the target so it survives visibly.
                r.underline = True
                paragraph.add_run(f" ({run['link']})").font.size = Pt(8)

    def add_list(block: dict, level: int = 0) -> None:
        style = "List Number" if block.get("ordered") else "List Bullet"
        # python-docx built-in list styles support levels 1-3 via suffixes.
        if level:
            style = f"{style} {min(level + 1, 3)}"
        for item in block.get("items") or []:
            paragraph = document.add_paragraph(style=style)
            if block.get("checklist"):
                paragraph.add_run("☑ " if item.get("checked") else "☐ ")
            add_runs(paragraph, item.get("runs") or [])
            for nested in item.get("children") or []:
                add_list(nested, level + 1)

    name_to_key = {a["name"]: a["key"] for a in (data.get("assets") or [])}
    for block in data.get("blocks") or []:
        btype = block.get("type")
        if btype == "heading":
            document.add_heading(
                _plain_text(block.get("runs") or []),
                level=min(int(block.get("level") or 1), 9),
            )
        elif btype == "quote":
            add_runs(
                document.add_paragraph(style="Intense Quote"), block.get("runs") or []
            )
        elif btype == "code":
            paragraph = document.add_paragraph()
            code_run = paragraph.add_run(block.get("text") or "")
            code_run.font.name = "Consolas"
            code_run.font.size = Pt(9)
        elif btype == "hr":
            separator = document.add_paragraph("· · ·")
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif btype == "list":
            add_list(block)
        elif btype == "table":
            rows = block.get("rows") or []
            if rows:
                width = max(len(r) for r in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_runs in enumerate(row):
                        cell_paragraph = table.cell(r_idx, c_idx).paragraphs[0]
                        add_runs(cell_paragraph, cell_runs)
        elif btype == "image":
            if block.get("asset"):
                key = name_to_key.get(block["asset"])
                if key:
                    # Editor resize is CSS px; Word wants inches (96 dpi).
                    # Cap at the ~6.5" content width of the default page.
                    width_px = block.get("width")
                    height_px = block.get("height")
                    if width_px:
                        size = {"width": Inches(min(width_px / 96, 6.5))}
                    elif height_px:
                        size = {"height": Inches(min(height_px / 96, 9))}
                    else:
                        size = {"width": Inches(6)}
                    try:
                        document.add_picture(io.BytesIO(read_blob(key)), **size)
                    except Exception:
                        # Unreadable/unsupported image: degrade to alt text
                        # rather than failing the whole document.
                        document.add_paragraph(block.get("alt") or "[image]")
            else:
                paragraph = document.add_paragraph(block.get("alt") or "")
                paragraph.add_run(f" ({block.get('url') or ''})")
        else:  # paragraph
            add_runs(document.add_paragraph(), block.get("runs") or [])

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _plain_text(runs: list[dict]) -> str:
    return "".join(run.get("text") or "" for run in runs)
