"""Export endpoints + worker: inline/job delivery, own-row isolation, and the
job-gated download path.

The download gate matters most here: an export artifact is a per-user
snapshot that may contain initiative-isolated content, so another guild
member must get 404 on the job and its download even though they share the
guild schema.
"""

from datetime import datetime, timedelta, timezone

import pytest
from httpx import AsyncClient

from app.core.config import settings
from app.models.platform.guild import GuildRole
from app.models.tenant.export_job import ExportJob, ExportJobStatus
from app.services.export import worker as export_worker
from app.services.storage import get_guild_storage
from app.testing.factories import create_task

pytestmark = pytest.mark.integration


@pytest.fixture(autouse=True)
def _tmp_uploads(monkeypatch, tmp_path):
    monkeypatch.setattr(settings, "UPLOADS_DIR", str(tmp_path))


async def _actor_with_tasks(acting_user, session, count=2):
    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    for i in range(count):
        await create_task(session, a.project, title=f"Task {i}")
    return a


async def test_inline_export_returns_pdf(client: AsyncClient, acting_user, session):
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    assert "attachment" in resp.headers["content-disposition"]
    assert resp.content.startswith(b"%PDF")


async def test_inline_export_csv_and_xlsx(client: AsyncClient, acting_user, session):
    a = await _actor_with_tasks(acting_user, session)

    csv_resp = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "csv"}
    )
    assert csv_resp.status_code == 200
    assert csv_resp.headers["content-type"].startswith("text/csv")
    body = csv_resp.content.decode("utf-8")
    assert "Task,Project,Status,Priority,Due,Assignees" in body
    assert "Task 0" in body and "Task 1" in body

    xlsx_resp = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "xlsx"}
    )
    assert xlsx_resp.status_code == 200
    assert xlsx_resp.content.startswith(b"PK")
    assert xlsx_resp.headers["content-type"].endswith("spreadsheetml.sheet")
    assert 'filename="tasks.xlsx"' in xlsx_resp.headers["content-disposition"]

    md_resp = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "md"}
    )
    assert md_resp.status_code == 200
    assert md_resp.headers["content-type"].startswith("text/markdown")
    md_body = md_resp.content.decode("utf-8")
    assert "| Task | Project | Status | Priority | Due | Assignees |" in md_body
    assert 'filename="tasks.md"' in md_resp.headers["content-disposition"]

    checklist = await client.get(
        a.g("/exports/tasks"),
        headers=a.headers,
        params={"format": "md", "layout": "checklist"},
    )
    assert checklist.status_code == 200
    body = checklist.content.decode("utf-8")
    assert "- [ ] Task 0" in body and "- [ ] Task 1" in body
    assert "| Task |" not in body  # checklist, not a table

    unknown = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "docx"}
    )
    assert unknown.status_code == 422  # HTTP-layer Literal validation


async def test_inline_export_respects_task_filters(
    client: AsyncClient, acting_user, session
):
    """Malformed conditions fail exactly like the list endpoint (same parse
    pipeline), proving the export rides the list's filter engine."""
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(
        a.g("/exports/tasks"),
        headers=a.headers,
        params={"conditions": "not json"},
    )
    assert resp.status_code == 400
    assert resp.json()["detail"] == "QUERY_INVALID_CONDITIONS"


async def test_export_max_rows_bound(
    client: AsyncClient, acting_user, session, monkeypatch
):
    monkeypatch.setattr(settings, "EXPORT_MAX_ROWS", 1)
    a = await _actor_with_tasks(acting_user, session, count=2)
    resp = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert resp.status_code == 400
    assert resp.json()["detail"] == "EXPORT_TOO_LARGE"


async def test_large_export_becomes_job(
    client: AsyncClient, acting_user, session, monkeypatch
):
    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", 0)
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert resp.status_code == 202
    body = resp.json()
    assert body["status"] == ExportJobStatus.queued.value
    assert body["source"] == "tasks"
    assert body["created_by_id"] == a.user.id
    # The row persists the SELECTOR only — no artifact_ref field is exposed,
    # and params echo the caller's own filter input.
    assert "artifact_ref" not in body
    assert body["params"]["include_archived"] is False

    # Not rendered yet: download must refuse, not serve a partial artifact.
    dl = await client.get(a.g(f"/exports/{body['id']}/download"), headers=a.headers)
    assert dl.status_code == 409
    assert dl.json()["detail"] == "EXPORT_NOT_READY"


async def test_job_limit_per_user(
    client: AsyncClient, acting_user, session, monkeypatch
):
    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", 0)
    monkeypatch.setattr(settings, "EXPORT_MAX_ACTIVE_JOBS_PER_USER", 1)
    a = await _actor_with_tasks(acting_user, session)
    first = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert first.status_code == 202
    second = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert second.status_code == 429
    assert second.json()["detail"] == "EXPORT_JOB_LIMIT_REACHED"


async def test_jobs_are_own_row_isolated(
    client: AsyncClient, acting_user, session, monkeypatch
):
    """Another member of the SAME guild sees neither the job nor its download
    (RLS hides the row -> 404); a guild admin sees it via the admin leg."""
    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", 0)
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert resp.status_code == 202
    job_id = resp.json()["id"]

    other = await acting_user(guild_role=GuildRole.member, guild=a.guild)
    admin = await acting_user(guild_role=GuildRole.admin, guild=a.guild)

    assert (
        await client.get(a.g(f"/exports/{job_id}"), headers=a.headers)
    ).status_code == 200
    for path in (f"/exports/{job_id}", f"/exports/{job_id}/download"):
        denied = await client.get(a.g(path), headers=other.headers)
        assert denied.status_code == 404, path
    assert (
        await client.get(a.g(f"/exports/{job_id}"), headers=admin.headers)
    ).status_code == 200

    # List views scope the same way.
    assert (await client.get(a.g("/exports/"), headers=other.headers)).json() == []
    assert [
        j["id"] for j in (await client.get(a.g("/exports/"), headers=a.headers)).json()
    ] == [job_id]


async def test_worker_renders_job_and_download_succeeds(
    client: AsyncClient, acting_user, session, monkeypatch, role_session
):
    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", 0)
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(a.g("/exports/tasks"), headers=a.headers)
    assert resp.status_code == 202
    job_id = resp.json()["id"]

    # The worker re-queries as the creator on an app_user session; point its
    # session factory at the test DB (the admin side is patched by the
    # standard harness).
    user_session = await role_session("app_user")
    monkeypatch.setattr(export_worker, "_open_user_session", lambda: user_session)

    await export_worker.process_export_jobs()

    status_resp = await client.get(a.g(f"/exports/{job_id}"), headers=a.headers)
    assert status_resp.status_code == 200
    body = status_resp.json()
    assert body["status"] == ExportJobStatus.done.value, body.get("error")
    assert body["expires_at"] is not None

    dl = await client.get(a.g(f"/exports/{job_id}/download"), headers=a.headers)
    assert dl.status_code == 200
    assert dl.headers["content-type"] == "application/pdf"
    assert dl.content.startswith(b"%PDF")

    # And the artifact never lands in the uploads table / media route.
    media = await client.get(
        f"/uploads/{a.guild.id}/exports/{job_id}.pdf", headers=a.headers
    )
    assert media.status_code == 404

    # The creator gets an inbox entry pointing at the finished job — the
    # recovery path when they navigated away while the render ran.
    from sqlmodel import select

    from app.models.platform.notification import Notification, NotificationType

    notifications = list(
        await session.exec(
            select(Notification).where(Notification.user_id == a.user.id)
        )
    )
    export_notes = [n for n in notifications if n.type == NotificationType.export_ready]
    assert len(export_notes) == 1
    assert export_notes[0].data["export_job_id"] == job_id
    assert export_notes[0].data["guild_id"] == a.guild.id


async def test_inline_project_export_returns_envelope(
    client: AsyncClient, acting_user, session
):
    """The engine-delivered project backup: same envelope the import endpoint
    consumes, same filename convention as the retired route."""
    import json

    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(
        a.g("/exports/project"),
        headers=a.headers,
        params={"project_id": a.project.id},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("application/json")
    assert ".initiative-project.json" in resp.headers["content-disposition"]
    envelope = json.loads(resp.content)
    assert envelope["schema_version"] >= 1
    assert envelope["project"]["name"] == a.project.name
    assert {t["title"] for t in envelope["tasks"]} == {"Task 0", "Task 1"}


async def test_project_export_report_formats(client: AsyncClient, acting_user, session):
    """pdf/csv/xlsx render the project report (unarchived tasks) from the same
    adapter that produces the json backup."""
    a = await _actor_with_tasks(acting_user, session)
    await create_task(session, a.project, title="Old news", is_archived=True)

    pdf = await client.get(
        a.g("/exports/project"),
        headers=a.headers,
        params={"project_id": a.project.id, "format": "pdf"},
    )
    assert pdf.status_code == 200
    assert pdf.content.startswith(b"%PDF")
    assert ".initiative-project" not in pdf.headers["content-disposition"]

    csv_resp = await client.get(
        a.g("/exports/project"),
        headers=a.headers,
        params={"project_id": a.project.id, "format": "csv"},
    )
    assert csv_resp.status_code == 200
    body = csv_resp.content.decode("utf-8")
    assert "Task,Status,Priority,Due,Assignees" in body
    assert "Task 0" in body and "Task 1" in body
    assert "Old news" not in body  # archived stays backup-only

    xlsx = await client.get(
        a.g("/exports/project"),
        headers=a.headers,
        params={"project_id": a.project.id, "format": "xlsx"},
    )
    assert xlsx.status_code == 200
    assert xlsx.content.startswith(b"PK")
    from io import BytesIO

    from openpyxl import load_workbook

    sheet = load_workbook(BytesIO(xlsx.content)).active
    cells = {cell.value for row in sheet.iter_rows() for cell in row}
    assert {"Task 0", "Task 1"} <= cells
    assert "Old news" not in cells  # archived exclusion holds in XLSX too


async def test_project_export_hidden_outside_initiative(
    client: AsyncClient, acting_user, session
):
    """A guild member outside the initiative gets 404 (RLS hides the project),
    exactly like the rest of the initiative boundary."""
    a = await _actor_with_tasks(acting_user, session)
    outsider = await acting_user(guild_role=GuildRole.member, guild=a.guild)
    resp = await client.get(
        a.g("/exports/project"),
        headers=outsider.headers,
        params={"project_id": a.project.id},
    )
    assert resp.status_code == 404


async def test_project_export_job_path_renders_json(
    client: AsyncClient, acting_user, session, monkeypatch, role_session
):
    import json

    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", 0)
    a = await _actor_with_tasks(acting_user, session)
    resp = await client.get(
        a.g("/exports/project"),
        headers=a.headers,
        params={"project_id": a.project.id},
    )
    assert resp.status_code == 202
    job_id = resp.json()["id"]
    assert resp.json()["source"] == "project"

    user_session = await role_session("app_user")
    monkeypatch.setattr(export_worker, "_open_user_session", lambda: user_session)
    await export_worker.process_export_jobs()

    dl = await client.get(a.g(f"/exports/{job_id}/download"), headers=a.headers)
    assert dl.status_code == 200, (
        (await client.get(a.g(f"/exports/{job_id}"), headers=a.headers))
        .json()
        .get("error")
    )
    assert dl.headers["content-type"].startswith("application/json")
    envelope = json.loads(dl.content)
    assert envelope["project"]["name"] == a.project.name


async def test_document_export_per_type_formats(
    client: AsyncClient, acting_user, session
):
    """Each document type exports only its own formats, with type-appropriate
    payloads; a mismatched combo is an immediate 400."""
    import json

    from app.models.tenant.document import DocumentType
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)

    async def export(doc, format):
        return await client.get(
            a.g("/exports/document"),
            headers=a.headers,
            params={"document_id": doc.id, "format": format},
        )

    # native (Lexical) -> the @lexical/file schema, named .lexical so the
    # editor toolbar's import button (which only accepts .lexical) takes it.
    native = await create_document(
        session,
        a.initiative,
        a.user,
        title="Notes",
        content={"root": {"children": [], "type": "root"}},
    )
    resp = await export(native, "json")
    assert resp.status_code == 200
    assert ".lexical" in resp.headers["content-disposition"]
    serialized = json.loads(resp.content)
    assert serialized["source"] == "Initiative"
    assert serialized["editorState"]["root"]["type"] == "root"
    assert isinstance(serialized["lastSaved"], int)

    # whiteboard -> standard Excalidraw file shape
    board = await create_document(
        session,
        a.initiative,
        a.user,
        title="Board",
        document_type=DocumentType.whiteboard,
        content={"elements": [{"type": "rectangle"}], "appState": {}, "files": {}},
    )
    resp = await export(board, "json")
    assert resp.status_code == 200
    scene = json.loads(resp.content)
    assert scene["type"] == "excalidraw"
    assert scene["elements"] == [{"type": "rectangle"}]

    # smart link -> markdown with the URL
    link = await create_document(
        session,
        a.initiative,
        a.user,
        title="Design doc",
        document_type=DocumentType.smart_link,
        content={"url": "https://example.com/spec"},
    )
    resp = await export(link, "md")
    assert resp.status_code == 200
    body = resp.content.decode("utf-8")
    assert "# Design doc" in body
    assert "<https://example.com/spec>" in body

    # mismatched combos: an immediate 400, no job side effects
    for doc, bad in ((native, "csv"), (board, "md"), (link, "json")):
        resp = await export(doc, bad)
        assert resp.status_code == 400, (doc.title, bad)
        assert resp.json()["detail"] == "EXPORT_INVALID_FORMAT"


async def test_document_export_lexical_formats(
    client: AsyncClient, acting_user, session
):
    """Lexical documents render to md/pdf/docx from one tree walk; a
    referenced same-guild image turns the markdown into a zip with assets and
    embeds into the pdf/docx."""
    import io
    import struct
    import zipfile
    import zlib

    from app.models.tenant.document import DocumentType  # noqa: F401
    from app.testing.factories import create_document

    def make_png():
        def chunk(tag, data):
            c = tag + data
            return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c))

        ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        return (
            b"\x89PNG\r\n\x1a\n"
            + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
            + chunk(b"IEND", b"")
        )

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    get_guild_storage(a.guild.id).write(
        "att-1.png", make_png(), content_type="image/png"
    )
    content = {
        "root": {
            "type": "root",
            "children": [
                {
                    "type": "heading",
                    "tag": "h2",
                    "children": [{"type": "text", "text": "Section", "format": 0}],
                },
                {
                    "type": "paragraph",
                    "children": [{"type": "text", "text": "Hello world", "format": 1}],
                },
                {
                    "type": "image",
                    "src": f"/uploads/{a.guild.id}/att-1.png",
                    "altText": "pic",
                },
            ],
        }
    }
    doc = await create_document(
        session, a.initiative, a.user, title="Rich Notes", content=content
    )

    async def export(format):
        return await client.get(
            a.g("/exports/document"),
            headers=a.headers,
            params={"document_id": doc.id, "format": format},
        )

    # md with an asset -> zip bundle
    md_resp = await export("md")
    assert md_resp.status_code == 200
    assert md_resp.headers["content-type"].startswith("application/zip")
    assert ".zip" in md_resp.headers["content-disposition"]
    archive = zipfile.ZipFile(io.BytesIO(md_resp.content))
    md_name = next(n for n in archive.namelist() if n.endswith(".md"))
    assert "assets/att-1.png" in archive.namelist()
    md_text = archive.read(md_name).decode("utf-8")
    assert "# Rich Notes" in md_text
    assert "**Hello world**" in md_text
    assert "![pic](assets/att-1.png)" in md_text

    # pdf embeds the staged image
    pdf_resp = await export("pdf")
    assert pdf_resp.status_code == 200
    assert pdf_resp.content.startswith(b"%PDF")

    # docx embeds the image too
    import docx

    docx_resp = await export("docx")
    assert docx_resp.status_code == 200
    assert docx_resp.headers["content-type"].endswith("wordprocessingml.document")
    document = docx.Document(io.BytesIO(docx_resp.content))
    assert any("Hello world" in p.text for p in document.paragraphs)
    assert len(document.inline_shapes) == 1


async def test_document_export_lexical_md_plain_without_assets(
    client: AsyncClient, acting_user, session
):
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    doc = await create_document(
        session,
        a.initiative,
        a.user,
        title="Plain",
        content={
            "root": {
                "type": "root",
                "children": [
                    {
                        "type": "paragraph",
                        "children": [{"type": "text", "text": "no images here"}],
                    }
                ],
            }
        },
    )
    resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": doc.id, "format": "md"},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("text/markdown")
    assert "no images here" in resp.content.decode("utf-8")


async def test_document_export_spreadsheet_formats(
    client: AsyncClient, acting_user, session
):
    from io import BytesIO

    from openpyxl import load_workbook

    from app.models.tenant.document import DocumentType
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    sheet_doc = await create_document(
        session,
        a.initiative,
        a.user,
        title="Budget: Q3",
        document_type=DocumentType.spreadsheet,
        content={
            "schema_version": 2,
            "dimensions": {"rows": 2, "cols": 2},
            # "=SUM(...)" is the user's own formula (first-class content);
            # "+not-a-formula" is a trigger-prefixed plain string (the CSV
            # smuggling vector the app never produces as a formula).
            "cells": {
                "0:0": "Item",
                "0:1": "=SUM(B2:B9)",
                "1:0": "+not-a-formula",
                "1:1": 42,
            },
            "cellStyles": {"0:0": {"bold": True, "fill": "#ff0000"}},
            "columns": {"0": {"width": 140}},
            "rows": {},
            "frozen": {"rows": 1, "cols": 0},
        },
    )

    csv_resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": sheet_doc.id, "format": "csv"},
    )
    assert csv_resp.status_code == 200
    text = csv_resp.content.decode("utf-8")
    assert "Item,=SUM(B2:B9)" in text  # own formulas survive
    assert "'+not-a-formula" in text  # non-formula triggers stay neutralized

    xlsx_resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": sheet_doc.id, "format": "xlsx"},
    )
    assert xlsx_resp.status_code == 200
    sheet = load_workbook(BytesIO(xlsx_resp.content)).active
    assert sheet.title == "Budget Q3"  # forbidden ":" stripped
    assert sheet.cell(row=1, column=1).value == "Item"
    assert sheet.cell(row=1, column=1).font.bold is True
    assert sheet.cell(row=1, column=2).value == "=SUM(B2:B9)"
    assert sheet.cell(row=1, column=2).data_type == "f"  # a live formula
    assert sheet.cell(row=2, column=1).value == "+not-a-formula"
    assert sheet.cell(row=2, column=1).data_type == "s"  # a string, not inferred
    assert sheet.cell(row=2, column=2).value == 42
    assert sheet.cell(row=2, column=2).data_type == "n"  # numbers stay typed
    assert sheet.freeze_panes == "A2"

    # json: the canonical snapshot in the importable envelope — content
    # round-trips verbatim (cells, styles, frozen panes, schema_version).
    import json

    json_resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": sheet_doc.id, "format": "json"},
    )
    assert json_resp.status_code == 200
    envelope = json.loads(json_resp.content)
    assert envelope["kind"] == "initiative-document"
    assert envelope["document_type"] == "spreadsheet"
    assert envelope["title"] == "Budget: Q3"
    assert envelope["content"] == sheet_doc.content


async def test_document_export_spreadsheet_survives_corrupt_snapshot(
    client: AsyncClient, acting_user, session
):
    """A renderer must not 500 over one bad snapshot entry: malformed cell
    keys are skipped, and 3-char / garbage hex colors are expanded / dropped
    instead of crashing openpyxl."""
    from io import BytesIO

    from openpyxl import load_workbook

    from app.models.tenant.document import DocumentType
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    doc = await create_document(
        session,
        a.initiative,
        a.user,
        title="Odd",
        document_type=DocumentType.spreadsheet,
        content={
            "schema_version": 2,
            "dimensions": {"rows": 1, "cols": 1},
            "cells": {"0:0": "ok", "corrupt": "x", "1:2:3": "y", ":": "z"},
            "cellStyles": {
                "0:0": {"fill": "#fff", "color": "not-a-color"},
            },
        },
    )
    resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": doc.id, "format": "xlsx"},
    )
    assert resp.status_code == 200
    sheet = load_workbook(BytesIO(resp.content)).active
    assert sheet.cell(row=1, column=1).value == "ok"
    # #fff shorthand expanded to a valid ARGB white fill.
    assert sheet.cell(row=1, column=1).fill.start_color.rgb == "FFFFFFFF"


async def test_document_export_file_passthrough(
    client: AsyncClient, acting_user, session, monkeypatch, role_session
):
    """File documents export the stored blob unconverted under the original
    name — inline and through the job path (job-id-prefixed artifact key)."""
    from app.models.tenant.document import DocumentType
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    payload = b"%PDF-original-bytes"
    get_guild_storage(a.guild.id).write(
        "stored-abc123.pdf", payload, content_type="application/pdf"
    )
    file_doc = await create_document(
        session,
        a.initiative,
        a.user,
        title="Uploaded report",
        document_type=DocumentType.file,
        file_url=f"/uploads/{a.guild.id}/stored-abc123.pdf",
        original_filename="Q3 Report Final.pdf",
        file_content_type="application/pdf",
        file_size=len(payload),
    )

    resp = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": file_doc.id, "format": "file"},
    )
    assert resp.status_code == 200
    assert resp.content == payload
    assert resp.headers["content-type"] == "application/pdf"
    # Spaces force the RFC 5987 form — the original name survives, escaped.
    assert "Q3%20Report%20Final.pdf" in resp.headers["content-disposition"]

    # Job path: the original filename survives via the job-id-prefixed key.
    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", -1)
    queued = await client.get(
        a.g("/exports/document"),
        headers=a.headers,
        params={"document_id": file_doc.id, "format": "file"},
    )
    assert queued.status_code == 202
    job_id = queued.json()["id"]

    user_session = await role_session("app_user")
    monkeypatch.setattr(export_worker, "_open_user_session", lambda: user_session)
    await export_worker.process_export_jobs()

    dl = await client.get(a.g(f"/exports/{job_id}/download"), headers=a.headers)
    assert dl.status_code == 200, (
        (await client.get(a.g(f"/exports/{job_id}"), headers=a.headers))
        .json()
        .get("error")
    )
    assert dl.content == payload
    assert "Q3%20Report%20Final.pdf" in dl.headers["content-disposition"]
    # The stored object carries the job id in its basename, so it can't
    # collide with another job exporting a same-named file (both backends
    # flatten a key to its basename, dropping any directory).
    from app.models.tenant.export_job import ExportJob

    job = await session.get(ExportJob, job_id)
    assert job.artifact_ref == f"exports/{job_id}-Q3 Report Final.pdf"


async def test_passthrough_exports_do_not_collide_by_filename(
    client: AsyncClient, acting_user, session, monkeypatch, role_session
):
    """Two members exporting a same-named file must not overwrite each other:
    the job id is in the storage basename (the directory a nested key would
    add is stripped by both backends), so each artifact is distinct."""
    from app.models.tenant.document import DocumentType
    from app.models.tenant.export_job import ExportJob
    from app.testing.factories import create_document

    monkeypatch.setattr(settings, "EXPORT_INLINE_MAX_ROWS", -1)
    a = await acting_user(guild_role=GuildRole.admin, initiative=True, project=True)
    storage = get_guild_storage(a.guild.id)

    async def queue_export(blob_key, blob_bytes):
        storage.write(blob_key, blob_bytes, content_type="application/pdf")
        doc = await create_document(
            session,
            a.initiative,
            a.user,
            title="report",
            document_type=DocumentType.file,
            file_url=f"/uploads/{a.guild.id}/{blob_key}",
            original_filename="report.pdf",  # SAME name for both
            file_content_type="application/pdf",
            file_size=len(blob_bytes),
        )
        resp = await client.get(
            a.g("/exports/document"),
            headers=a.headers,
            params={"document_id": doc.id, "format": "file"},
        )
        assert resp.status_code == 202
        return resp.json()["id"]

    job_a = await queue_export("src-a.pdf", b"AAAA-first-member")
    job_b = await queue_export("src-b.pdf", b"BBBB-second-member")

    user_session = await role_session("app_user")
    monkeypatch.setattr(export_worker, "_open_user_session", lambda: user_session)
    await export_worker.process_export_jobs()

    row_a = await session.get(ExportJob, job_a)
    row_b = await session.get(ExportJob, job_b)
    assert row_a.artifact_ref != row_b.artifact_ref

    dl_a = await client.get(a.g(f"/exports/{job_a}/download"), headers=a.headers)
    dl_b = await client.get(a.g(f"/exports/{job_b}/download"), headers=a.headers)
    assert dl_a.content == b"AAAA-first-member"  # not overwritten by job_b
    assert dl_b.content == b"BBBB-second-member"


async def test_document_export_hidden_outside_initiative(
    client: AsyncClient, acting_user, session
):
    from app.testing.factories import create_document

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    doc = await create_document(session, a.initiative, a.user, title="Secret")
    outsider = await acting_user(guild_role=GuildRole.member, guild=a.guild)
    resp = await client.get(
        a.g("/exports/document"),
        headers=outsider.headers,
        params={"document_id": doc.id, "format": "json"},
    )
    assert resp.status_code == 404


async def test_gc_expires_artifacts(acting_user, session):
    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    storage = get_guild_storage(a.guild.id)
    key = "exports/424242.pdf"
    storage.write(key, b"%PDF-fake", content_type="application/pdf")
    job = ExportJob(
        guild_id=a.guild.id,
        created_by_id=a.user.id,
        source="tasks",
        template_id="task-table",
        format="pdf",
        status=ExportJobStatus.done,
        artifact_ref=key,
        expires_at=datetime.now(timezone.utc) - timedelta(hours=1),
    )
    session.add(job)
    await session.commit()

    await export_worker.process_export_gc()

    assert storage.open_readable(key) is None
    from app.testing import route_session_to_guild

    session.expunge_all()
    await route_session_to_guild(session, a.guild.id)
    refreshed = await session.get(ExportJob, job.id)
    assert refreshed.status == ExportJobStatus.expired.value
    assert refreshed.artifact_ref is None


async def test_gc_expires_row_even_when_artifact_delete_fails(
    acting_user, session, monkeypatch
):
    """A failing storage delete must not pin the job in ``done`` (it would
    re-fail every pass and abort GC for every guild after it) — the row still
    expires and the failure is only logged."""
    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    job = ExportJob(
        guild_id=a.guild.id,
        created_by_id=a.user.id,
        source="tasks",
        template_id="task-table",
        format="pdf",
        status=ExportJobStatus.done,
        artifact_ref="exports/31337.pdf",
        expires_at=datetime.now(timezone.utc) - timedelta(hours=1),
    )
    session.add(job)
    await session.commit()

    class _BrokenStorage:
        def delete(self, key: str) -> bool:
            raise RuntimeError("storage down")

    import app.services.storage as storage_module

    monkeypatch.setattr(
        storage_module, "get_guild_storage", lambda gid: _BrokenStorage()
    )

    await export_worker.process_export_gc()

    from app.testing import route_session_to_guild

    session.expunge_all()
    await route_session_to_guild(session, a.guild.id)
    refreshed = await session.get(ExportJob, job.id)
    assert refreshed.status == ExportJobStatus.expired.value
    assert refreshed.artifact_ref is None


async def _queue_with_items(acting_user, session):
    from app.models.tenant.queue import QueueItemDocument, QueueItemTag, QueueItemTask
    from app.testing.factories import (
        create_document,
        create_queue,
        create_queue_item,
        create_tag,
    )

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    queue = await create_queue(
        session,
        a.initiative,
        a.user,
        name="Battle Order",
        description="Round tracker",
        is_active=True,
        current_round=3,
    )
    # Rotation order is position DESCENDING: Alice acts first.
    current = await create_queue_item(session, queue, label="Alice", position=30)
    await create_queue_item(
        session, queue, label="Bob", position=20, held_at_round=2, notes="waiting"
    )
    lurker = await create_queue_item(
        session, queue, label="Lurker", position=10, is_visible=False
    )
    tag = await create_tag(session, a.guild, name="npc")
    session.add(QueueItemTag(queue_item_id=lurker.id, tag_id=tag.id))
    doc = await create_document(session, a.initiative, a.user, title="Dungeon map")
    task = await create_task(session, a.project, title="Prep loot")
    session.add(
        QueueItemDocument(
            queue_item_id=current.id, document_id=doc.id, guild_id=a.guild.id
        )
    )
    session.add(
        QueueItemTask(queue_item_id=current.id, task_id=task.id, guild_id=a.guild.id)
    )
    queue.current_item_id = current.id
    session.add(queue)
    await session.commit()
    return a, queue


async def test_queue_export_json_envelope(client: AsyncClient, acting_user, session):
    import json

    a, queue = await _queue_with_items(acting_user, session)
    resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "json"},
    )
    assert resp.status_code == 200
    assert ".initiative-queue.json" in resp.headers["content-disposition"]
    envelope = json.loads(resp.content)
    assert envelope["kind"] == "initiative-queue"
    assert envelope["schema_version"] == 1
    assert envelope["name"] == "Battle Order"
    assert envelope["is_active"] is True
    assert envelope["current_round"] == 3

    items = envelope["items"]
    assert [i["label"] for i in items] == ["Alice", "Bob", "Lurker"]  # rotation order
    assert [i["is_current"] for i in items] == [True, False, False]
    assert items[1]["held_at_round"] == 2
    assert items[2]["is_visible"] is False
    assert items[2]["tags"] == ["npc"]
    # Guild-local references ride along as display text only (an import can't
    # rebind them): member name, linked document/task titles — never raw ids.
    assert "member" in items[0] and "user_id" not in items[0]
    assert items[0]["documents"] == ["Dungeon map"]
    assert items[0]["tasks"] == ["Prep loot"]
    assert items[1]["documents"] == [] and items[1]["tasks"] == []


async def test_queue_export_report_formats(client: AsyncClient, acting_user, session):
    import io

    from pypdf import PdfReader

    a, queue = await _queue_with_items(acting_user, session)

    csv_resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "csv"},
    )
    assert csv_resp.status_code == 200
    text = csv_resp.content.decode("utf-8")
    assert "#,Item,Member,Tags,Notes,Status" in text
    assert "1,Alice,,,,Current" in text
    assert "2,Bob,,,waiting,Held" in text
    assert "3,Lurker,,npc,,Hidden" in text

    md_resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "md"},
    )
    assert md_resp.status_code == 200
    md = md_resp.content.decode("utf-8")
    assert "# Battle Order" in md
    assert "1. **Alice** (Current)" in md  # numbered turn order, current bolded
    assert "2. Bob (waiting · Held)" in md
    assert "3. Lurker (npc · Hidden)" in md
    assert "| Item |" not in md  # a list, not a table

    pdf_resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "pdf"},
    )
    assert pdf_resp.status_code == 200
    assert pdf_resp.content.startswith(b"%PDF")
    pdf_text = PdfReader(io.BytesIO(pdf_resp.content)).pages[0].extract_text()
    assert "Battle Order" in pdf_text
    assert "Alice" in pdf_text and "Lurker" in pdf_text
    assert "Round tracker" in pdf_text  # description block rendered


async def _counter_group_with_counters(acting_user, session):
    from decimal import Decimal

    from app.testing.factories import create_counter, create_counter_group

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    group = await create_counter_group(
        session, a.initiative, a.user, name="Party Resources", description="Session 12"
    )
    await create_counter(
        session,
        group,
        name="Torches",
        count=Decimal("5.0000000000"),
        min=Decimal("0"),
        max=Decimal("10"),
        position=Decimal("1"),
    )
    await create_counter(
        session,
        group,
        name="Rations",
        count=Decimal("2.5"),
        step=Decimal("0.5"),
        position=Decimal("2"),
    )
    return a, group


async def test_counter_group_export_json_envelope(
    client: AsyncClient, acting_user, session
):
    import json

    a, group = await _counter_group_with_counters(acting_user, session)
    resp = await client.get(
        a.g("/exports/counter-group"),
        headers=a.headers,
        params={"counter_group_id": group.id, "format": "json"},
    )
    assert resp.status_code == 200
    assert ".initiative-counter-group.json" in resp.headers["content-disposition"]
    envelope = json.loads(resp.content)
    assert envelope["kind"] == "initiative-counter-group"
    assert envelope["schema_version"] == 1
    assert envelope["name"] == "Party Resources"

    by_name = {c["name"]: c for c in envelope["counters"]}
    # NUMERIC(20,10) scale must not leak: integral Decimals become plain ints.
    assert by_name["Torches"]["count"] == 5
    assert isinstance(by_name["Torches"]["count"], int)
    assert by_name["Torches"]["min"] == 0 and by_name["Torches"]["max"] == 10
    assert by_name["Rations"]["count"] == 2.5
    assert by_name["Rations"]["step"] == 0.5
    assert by_name["Rations"]["min"] is None
    assert by_name["Torches"]["view_mode"] == "number"


async def test_counter_group_export_report_formats(
    client: AsyncClient, acting_user, session
):
    import io
    from io import BytesIO

    from openpyxl import load_workbook
    from pypdf import PdfReader

    a, group = await _counter_group_with_counters(acting_user, session)

    csv_resp = await client.get(
        a.g("/exports/counter-group"),
        headers=a.headers,
        params={"counter_group_id": group.id, "format": "csv"},
    )
    assert csv_resp.status_code == 200
    text = csv_resp.content.decode("utf-8")
    assert "Counter,Count,Min,Max,Step" in text
    assert "Torches,5,0,10,1" in text
    assert "Rations,2.5,,,0.5" in text  # unbounded min/max stay empty

    xlsx_resp = await client.get(
        a.g("/exports/counter-group"),
        headers=a.headers,
        params={"counter_group_id": group.id, "format": "xlsx"},
    )
    assert xlsx_resp.status_code == 200
    sheet = load_workbook(BytesIO(xlsx_resp.content)).active
    assert [c.value for c in sheet[1]] == ["Counter", "Count", "Min", "Max", "Step"]
    assert sheet.cell(row=2, column=2).value == 5
    assert sheet.cell(row=2, column=2).data_type == "n"  # counts stay numeric
    assert sheet.cell(row=3, column=2).value == 2.5

    pdf_resp = await client.get(
        a.g("/exports/counter-group"),
        headers=a.headers,
        params={"counter_group_id": group.id, "format": "pdf"},
    )
    assert pdf_resp.status_code == 200
    assert pdf_resp.content.startswith(b"%PDF")
    pdf_text = PdfReader(io.BytesIO(pdf_resp.content)).pages[0].extract_text()
    assert "Party Resources" in pdf_text
    assert "Torches" in pdf_text and "Rations" in pdf_text


async def test_tool_exports_hidden_outside_initiative(
    client: AsyncClient, acting_user, session
):
    """The initiative gate: a same-guild member outside the initiative gets
    404 (RLS hides the row) for both tool sources."""
    a, queue = await _queue_with_items(acting_user, session)
    outsider = await acting_user(guild_role=GuildRole.member, guild=a.guild)

    resp = await client.get(
        a.g("/exports/queue"),
        headers=outsider.headers,
        params={"queue_id": queue.id, "format": "json"},
    )
    assert resp.status_code == 404

    from app.testing.factories import create_counter_group

    group = await create_counter_group(
        session, a.initiative, a.user, name="Secret counters"
    )
    resp = await client.get(
        a.g("/exports/counter-group"),
        headers=outsider.headers,
        params={"counter_group_id": group.id, "format": "json"},
    )
    assert resp.status_code == 404


async def _set_locale(session, user, locale):
    """Persist a locale on the creator so the request (and the worker replay)
    re-load it — the export content localizes to the creator, not the
    request's Accept-Language."""
    db_user = await session.get(type(user), user.id)
    db_user.locale = locale
    session.add(db_user)
    await session.commit()


async def test_task_export_localizes_report_content(
    client: AsyncClient, acting_user, session
):
    """The generated report chrome (column headers, title, summary line)
    localizes to the creator's locale — not just the surrounding UI."""
    a = await _actor_with_tasks(acting_user, session, count=1)
    await _set_locale(session, a.user, "es")

    csv_resp = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "csv"}
    )
    assert csv_resp.status_code == 200
    body = csv_resp.content.decode("utf-8")
    # Spanish headers, not "Task,Project,Status,...".
    assert "Tarea,Proyecto,Estado,Prioridad,Vence,Asignados" in body

    md_resp = await client.get(
        a.g("/exports/tasks"), headers=a.headers, params={"format": "md"}
    )
    md = md_resp.content.decode("utf-8")
    assert "# Tareas" in md  # localized title
    assert "1 tarea · generado el" in md  # localized, singular plural form


async def test_queue_export_localizes_status_flags_and_headers(
    client: AsyncClient, acting_user, session
):
    """Queue report status flags (Current/Held/Hidden) and headers localize —
    the JSON envelope stays canonical (importable machine data)."""
    import json

    a, queue = await _queue_with_items(acting_user, session)
    await _set_locale(session, a.user, "fr")

    csv_resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "csv"},
    )
    body = csv_resp.content.decode("utf-8")
    assert "N°,Élément,Membre,Étiquettes,Notes,Statut" in body  # French headers
    assert "Actuel" in body and "En attente" in body and "Masqué" in body  # flags

    # The importable envelope must NOT be localized — kind and field keys stay
    # canonical so a French user's backup still imports.
    json_resp = await client.get(
        a.g("/exports/queue"),
        headers=a.headers,
        params={"queue_id": queue.id, "format": "json"},
    )
    envelope = json.loads(json_resp.content)
    assert envelope["kind"] == "initiative-queue"
    assert "items" in envelope and "is_current" in envelope["items"][0]


async def test_task_detailed_pdf_is_one_page_per_task_with_full_detail(
    client: AsyncClient, acting_user, session
):
    """layout=detailed renders a one-task-per-page PDF carrying each task's
    description, subtasks and comments — not the tabular line-per-task list."""
    import io

    from pypdf import PdfReader

    from app.testing.factories import create_comment, create_subtask

    a = await acting_user(guild_role=GuildRole.member, initiative=True, project=True)
    t1 = await create_task(
        session,
        a.project,
        title="Boss fight",
        description="Balance the encounter.\nCheck the second phase.",
    )
    await create_subtask(session, t1, content="Tune the HP", is_completed=True)
    await create_subtask(session, t1, content="Write the dialogue")
    await create_comment(session, a.user, task=t1, content="Started already.")
    await create_task(session, a.project, title="Loot table")

    resp = await client.get(
        a.g("/exports/tasks"),
        headers=a.headers,
        params={"format": "pdf", "layout": "detailed"},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"] == "application/pdf"
    reader = PdfReader(io.BytesIO(resp.content))
    assert len(reader.pages) == 2  # one page per task
    text = "\n".join(p.extract_text() for p in reader.pages)
    assert "Boss fight" in text and "Loot table" in text
    assert "Balance the encounter" in text  # description
    assert "Check the second phase" in text  # description line break preserved
    assert "Tune the HP" in text and "Write the dialogue" in text  # subtasks
    assert "Started already" in text  # comment body
    # Localized section labels (en locale).
    for label in ("Description", "Subtasks", "Comments"):
        assert label in text


async def test_detailed_layout_ignored_for_non_pdf_formats(
    client: AsyncClient, acting_user, session
):
    """layout=detailed only applies to PDF; csv falls through to the table."""
    a = await _actor_with_tasks(acting_user, session, count=1)
    resp = await client.get(
        a.g("/exports/tasks"),
        headers=a.headers,
        params={"format": "csv", "layout": "detailed"},
    )
    assert resp.status_code == 200
    assert resp.headers["content-type"].startswith("text/csv")
    assert "Task,Project,Status,Priority,Due,Assignees" in resp.content.decode("utf-8")
